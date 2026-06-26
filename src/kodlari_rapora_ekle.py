from docx import Document

report_path = "report/HIGGS_Final_Raporu.docx"

code_files = [
    "src/rastgele_orneklem.py",
    "src/main.py",
    "src/modelleme.py",
    "src/grafikler.py",
    "src/roc_grafigi.py",
    "src/roc_ova.py",
    "src/flowchart_olustur.py",
    "src/rapor_olustur.py",
]

doc = Document(report_path)

doc.add_page_break()
doc.add_heading("12. Ek: Projede Kullanilan Kodlar", level=1)

for file_path in code_files:
    doc.add_heading(file_path, level=2)

    with open(file_path, "r", encoding="utf-8") as f:
        code = f.read()

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = None

doc.save(report_path)

print("Kodlar rapora eklendi:", report_path)