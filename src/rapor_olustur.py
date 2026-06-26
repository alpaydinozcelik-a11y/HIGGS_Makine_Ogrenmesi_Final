import os
import pandas as pd
from docx import Document
from docx.shared import Inches

os.makedirs("report", exist_ok=True)

doc = Document()

doc.add_heading("HIGGS Dataset Makine Ogrenmesi Final Projesi", 0)

doc.add_heading("1. Projenin Amaci", level=1)
doc.add_paragraph(
    "Bu projede HIGGS veri seti uzerinde makine ogrenmesi pipeline'i olusturulmustur. "
    "Calismada veri on isleme, aykiri deger analizi, ozellik secimi, nested cross-validation "
    "ile hiperparametre optimizasyonu ve model performans karsilastirmasi yapilmistir."
)

doc.add_heading("2. Veri Seti", level=1)
doc.add_paragraph(
    "Projede UCI Machine Learning Repository uzerinde bulunan HIGGS veri seti kullanilmistir. "
    "Veri seti Higgs bozonu sinyali ile arka plan olaylarini ayirt etmeyi amaclayan ikili "
    "siniflandirma problemidir. Orijinal veri seti 11 milyon ornekten ve 28 sayisal ozellikten "
    "olusur. Bu projede tum veri seti taranarak rastgele 100.000 ornek secilmistir."
)

doc.add_heading("3. Veri On Isleme", level=1)
doc.add_paragraph(
    "Aykiri deger analizi icin IQR yontemi kullanilmistir. Her ozellik icin Q1, Q3 ve IQR "
    "degerleri hesaplanmis; alt sinir Q1 - 1.5*IQR, ust sinir Q3 + 1.5*IQR olarak belirlenmistir. "
    "Sinir disinda kalan degerler silinmemis, sinir degerlerine cekilmistir."
)
doc.add_paragraph(
    "Aykiri deger isleminden sonra tum sayisal ozellikler MinMaxScaler ile [0, 1] araligina "
    "donusturulmustur."
)

doc.add_heading("4. Ozellik Secimi", level=1)
doc.add_paragraph(
    "Ozellik secimi asamasinda filtre tabanli yontemlerden ANOVA F-score kullanilmistir. "
    "Proje kapsaminda en iyi 15 ozellik secilmistir."
)

selected_features = pd.read_csv("data/selected_features.csv")

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = "Ozellik"
table.rows[0].cells[1].text = "Skor"

for _, row in selected_features.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(row["feature"])
    cells[1].text = f"{row['score']:.4f}"

doc.add_paragraph("Sekil 1. ANOVA F-score ile secilen en iyi 15 ozellik.")
doc.add_picture("figures/selected_features.png", width=Inches(6))

doc.add_heading("5. Nested Cross-Validation Akislari", level=1)
doc.add_paragraph(
    "Nested cross-validation yapisinda dis dongu modelin genelleme performansini olcmek icin, "
    "ic dongu ise ozellik secimi ve hiperparametre optimizasyonu icin kullanilmistir."
)

doc.add_paragraph("Sekil 2. Flowchart A: Ic dongude ozellik secimi.")
doc.add_picture("figures/flowchart_feature_selection.png", width=Inches(6))

doc.add_paragraph("Sekil 3. Flowchart B: Ic dongude hiperparametre aramasi.")
doc.add_picture("figures/flowchart_hyperparameter.png", width=Inches(6))

doc.add_heading("6. Modelleme ve Degerlendirme", level=1)
doc.add_paragraph(
    "Model degerlendirme asamasinda nested cross-validation yontemi kullanilmistir. "
    "Dis dongu 5-fold, ic dongu 3-fold olarak tasarlanmistir. Ic dongude GridSearchCV ile "
    "hiperparametre optimizasyonu yapilmis, dis dongude test performansi degerlendirilmistir."
)
doc.add_paragraph(
    "Kullanilan modeller: KNN, SVM, MLP ve XGBoost."
)

doc.add_heading("7. Performans Sonuclari", level=1)

summary_df = pd.read_csv("data/model_summary.csv")
if "model" not in summary_df.columns:
    summary_df = summary_df.rename(columns={summary_df.columns[0]: "model"})

table = doc.add_table(rows=1, cols=6)
table.style = "Table Grid"
headers = ["Model", "Accuracy", "Precision", "Recall", "F1 Score", "ROC-AUC"]

for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

for _, row in summary_df.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(row["model"])
    cells[1].text = f"{row['accuracy']:.4f}"
    cells[2].text = f"{row['precision']:.4f}"
    cells[3].text = f"{row['recall']:.4f}"
    cells[4].text = f"{row['f1_score']:.4f}"
    cells[5].text = f"{row['roc_auc']:.4f}"

doc.add_paragraph("Sekil 4. Modellerin ortalama performans karsilastirmasi.")
doc.add_picture("figures/model_performance.png", width=Inches(6))

doc.add_heading("8. ROC Egrileri", level=1)
doc.add_paragraph(
    "ROC egrileri modellerin farkli esik degerlerinde siniflari ayirt etme basarisini gosterir. "
    "Binary siniflandirma problemi icin OVA yaklasimiyla hem Class 1 hem de Class 0 icin ROC "
    "egrileri cizilmistir."
)

doc.add_paragraph("Sekil 5. Modellerin genel ROC egrileri.")
doc.add_picture("figures/roc_curves.png", width=Inches(6))

doc.add_paragraph("Sekil 6. OVA ROC egrileri - Class 1.")
doc.add_picture("figures/roc_ova_class1.png", width=Inches(6))

doc.add_paragraph("Sekil 7. OVA ROC egrileri - Class 0.")
doc.add_picture("figures/roc_ova_class0.png", width=Inches(6))

doc.add_heading("9. En Basarili Model Yorumu", level=1)

best_model = summary_df.sort_values("roc_auc", ascending=False).iloc[0]

doc.add_paragraph(
    f"Elde edilen sonuclara gore en yuksek ROC-AUC degerine sahip model {best_model['model']} olmustur. "
    f"{best_model['model']} modeli {best_model['roc_auc']:.4f} ROC-AUC, "
    f"{best_model['f1_score']:.4f} F1-score ve {best_model['accuracy']:.4f} accuracy degeri elde etmistir. "
    "Bu nedenle HIGGS veri setinde siniflari ayirt etmede en basarili model olarak degerlendirilmistir."
)

doc.add_heading("10. Sonuc", level=1)
doc.add_paragraph(
    "Bu calismada HIGGS veri seti uzerinde uctan uca bir makine ogrenmesi sureci uygulanmistir. "
    "Sonuclar, XGBoost modelinin diger modellere gore daha yuksek ROC-AUC performansi verdigini gostermektedir."
)

doc.add_heading("11. Kaynakca", level=1)
doc.add_paragraph(
    "Whiteson, D. (2014). HIGGS [Dataset]. UCI Machine Learning Repository. "
    "https://doi.org/10.24432/C5V312"
)

doc.save("report/HIGGS_Final_Raporu.docx")

print("Rapor olusturuldu: report/HIGGS_Final_Raporu.docx")